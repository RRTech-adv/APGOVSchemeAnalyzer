import os
from typing import Optional
from pathlib import Path

try:
    import pdfplumber
    PDF_PARSER_AVAILABLE = True
except ImportError:
    PDF_PARSER_AVAILABLE = False
    try:
        import fitz  # PyMuPDF
        PDF_PARSER_AVAILABLE = True
    except ImportError:
        pass

try:
    from docx import Document
    DOCX_PARSER_AVAILABLE = True
except ImportError:
    DOCX_PARSER_AVAILABLE = False

class ParserService:
    """Service for parsing different file formats"""
    
    def __init__(self):
        self.pdf_available = PDF_PARSER_AVAILABLE
        self.docx_available = DOCX_PARSER_AVAILABLE
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract text from a file based on its extension
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text or None if parsing fails
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == ".pdf":
            return self._extract_from_pdf(file_path)
        elif file_ext == ".docx":
            return self._extract_from_docx(file_path)
        elif file_ext == ".txt":
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file"""
        if not self.pdf_available:
            raise ImportError("No PDF parser available. Install pdfplumber or PyMuPDF (fitz)")
        
        text_content = []
        
        # Try pdfplumber first (better for complex layouts)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                print(f"Extracting text from PDF with {total_pages} pages...")
                
                for i, page in enumerate(pdf.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                        
                        # Progress indicator for large PDFs
                        if (i + 1) % 10 == 0 or (i + 1) == total_pages:
                            print(f"Processed {i + 1}/{total_pages} pages...")
                    except Exception as page_error:
                        print(f"Warning: Error extracting text from page {i + 1}: {page_error}")
                        continue
                
                print(f"PDF extraction complete. Extracted {len(''.join(text_content))} characters.")
                return "\n".join(text_content)
        except ImportError:
            pass
        
        # Fallback to PyMuPDF (faster for large files)
        try:
            import fitz
            doc = fitz.open(file_path)
            total_pages = len(doc)
            print(f"Extracting text from PDF with {total_pages} pages using PyMuPDF...")
            
            for page_num in range(total_pages):
                try:
                    page = doc[page_num]
                    text = page.get_text()
                    if text:
                        text_content.append(text)
                    
                    # Progress indicator for large PDFs
                    if (page_num + 1) % 10 == 0 or (page_num + 1) == total_pages:
                        print(f"Processed {page_num + 1}/{total_pages} pages...")
                except Exception as page_error:
                    print(f"Warning: Error extracting text from page {page_num + 1}: {page_error}")
                    continue
            
            doc.close()
            print(f"PDF extraction complete. Extracted {len(''.join(text_content))} characters.")
            return "\n".join(text_content)
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return None
    
    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file"""
        if not self.docx_available:
            raise ImportError("python-docx not installed. Install it using: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_content = []
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            return "\n".join(text_content)
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return None
    
    def _extract_from_txt(self, file_path: str) -> Optional[str]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT file: {e}")
            return None

